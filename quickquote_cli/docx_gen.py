"""Generate CES proposal .docx files from the template using python-docx.

The template contains placeholder tokens like {{PROJECT_NAME}} and a sample
first bid-item section. generate_proposal() substitutes the tokens, writes
the correct billing-mode fee sentence, and clones paragraphs for additional
sections.
"""
import os
import re
from copy import deepcopy

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .filename import safe_name, unique_path
from .formatting import build_fee_text, format_fee_for_doc
from .paths import PLACEHOLDERS, TEMPLATE_CONSULTING

_XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"

# Lightweight markdown markers for list rendering. Stripped from the line
# text and converted to Word's built-in "List Bullet" / "List Number" styles
# when present. Mirrors src/components/DocPreview.tsx parseRichText.
_BULLET_RE = re.compile(r"^(\s*)[-*]\s+(.*)$")
_NUMBER_RE = re.compile(r"^(\s*)\d+\.\s+(.*)$")


# ── docx paragraph/run helpers ────────────────────────────────────────────────

def _set_run_with_linebreaks(run, text: str) -> None:
    r = run._r
    for t in r.findall(qn("w:t")):
        r.remove(t)
    for br in r.findall(qn("w:br")):
        r.remove(br)
    parts = text.split("\n")
    for i, part in enumerate(parts):
        if i > 0:
            r.append(OxmlElement("w:br"))
        t = OxmlElement("w:t")
        t.text = part
        if part != part.strip():
            t.set(_XML_SPACE, "preserve")
        r.append(t)


def _replace_in_paragraph(paragraph, old: str, new: str) -> None:
    if old not in paragraph.text:
        return
    for run in paragraph.runs:
        if old in run.text:
            replacement = run.text.replace(old, new)
            if "\n" in replacement:
                _set_run_with_linebreaks(run, replacement)
            else:
                run.text = replacement
            return
    # Fallback: placeholder split across runs.
    full    = paragraph.text
    rebuilt = full.replace(old, new)
    if paragraph.runs:
        if "\n" in rebuilt:
            _set_run_with_linebreaks(paragraph.runs[0], rebuilt)
        else:
            paragraph.runs[0].text = rebuilt
        for run in paragraph.runs[1:]:
            run.text = ""


def _replace_in_header_textboxes(hdr_element, replacements: dict) -> None:
    """Replace placeholders inside text boxes in a header/footer.

    Text box paragraphs are nested inside w:txbxContent elements and are not
    exposed through python-docx's .paragraphs property. Placeholders may also
    be split across multiple runs, so we merge all w:t text per paragraph
    before substituting.
    """
    for txbx_content in hdr_element.iter(qn("w:txbxContent")):
        for p_el in txbx_content.findall(qn("w:p")):
            runs  = p_el.findall(".//" + qn("w:r"))
            all_t = [t for r in runs for t in r.findall(qn("w:t"))]
            if not all_t:
                continue
            full = "".join(t.text or "" for t in all_t)
            new  = full
            for old, val in replacements.items():
                new = new.replace(old, val)
            if new == full:
                continue
            all_t[0].text = new
            if new != new.strip():
                all_t[0].set(_XML_SPACE, "preserve")
            for t in all_t[1:]:
                t.text = ""


def _set_paragraph_style(p_el, style_name: str) -> None:
    """Force-set the Word paragraph style on a w:p element.

    Replaces any existing w:pStyle inside w:pPr; creates the pPr/pStyle
    elements if missing. Used to flip cloned scope paragraphs to "List
    Bullet" / "List Number" when the source line carried a markdown marker.
    """
    pPr = p_el.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_el.insert(0, pPr)
    existing = pPr.find(qn("w:pStyle"))
    if existing is not None:
        pPr.remove(existing)
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), style_name)
    pPr.insert(0, pStyle)


def _insert_scope_paragraphs(anchor_para, title_para, entries: list[tuple[str, object]]) -> None:
    """Insert real Word paragraphs after *anchor_para*.

    entries: list of (kind, value):
      ("title", text)   -- bold paragraph copying title_para formatting
      ("scope", text)   -- plain paragraph copying anchor_para formatting.
                           Lines starting with "- " / "* " / "<n>. " switch
                           to the matching List Bullet / List Number style;
                           the marker is stripped from the run text.
      ("xml",   element)-- pre-built lxml element inserted as-is
    """
    parent    = anchor_para._p.getparent()
    insert_at = list(parent).index(anchor_para._p) + 1

    for i, (kind, value) in enumerate(entries):
        if kind == "xml":
            new_p = value
        elif kind == "title":
            new_p = OxmlElement("w:p")
            orig_ppr = title_para._p.find(qn("w:pPr"))
            if orig_ppr is not None:
                new_p.append(deepcopy(orig_ppr))
            r = OxmlElement("w:r")
            orig_runs = title_para._p.findall(qn("w:r"))
            if orig_runs:
                orig_rpr = orig_runs[0].find(qn("w:rPr"))
                rpr = deepcopy(orig_rpr) if orig_rpr is not None else OxmlElement("w:rPr")
            else:
                rpr = OxmlElement("w:rPr")
            if rpr.find(qn("w:b")) is None:
                rpr.insert(0, OxmlElement("w:b"))
            r.append(rpr)
            t = OxmlElement("w:t")
            t.text = value
            r.append(t)
            new_p.append(r)
        else:  # "scope"
            list_style = None
            text = value if isinstance(value, str) else ""
            if isinstance(value, str):
                m = _BULLET_RE.match(value)
                if m:
                    list_style = "List Bullet"
                    text = m.group(2)
                else:
                    m = _NUMBER_RE.match(value)
                    if m:
                        list_style = "List Number"
                        text = m.group(2)

            new_p = deepcopy(anchor_para._p)
            for r_el in new_p.findall(qn("w:r")):
                new_p.remove(r_el)
            if list_style:
                _set_paragraph_style(new_p, list_style)
            if text:
                r = OxmlElement("w:r")
                orig_runs = anchor_para._p.findall(qn("w:r"))
                if orig_runs:
                    orig_rpr = orig_runs[0].find(qn("w:rPr"))
                    if orig_rpr is not None:
                        r.append(deepcopy(orig_rpr))
                t = OxmlElement("w:t")
                t.text = text
                if text != text.strip():
                    t.set(_XML_SPACE, "preserve")
                r.append(t)
                new_p.append(r)

        parent.insert(insert_at + i, new_p)


def _exclusions_entries(exclusions: str) -> list[tuple[str, object]]:
    """Build paragraph entries for an Exclusions block.

    Renders as a bold "Scope specifically excluded:" heading followed by the
    exclusions content, line-by-line. The line-by-line emission lets the
    scope renderer detect markdown list markers and apply List Bullet /
    List Number styles, the same way scope-of-work bullets work.
    Empty/whitespace-only input emits nothing.
    """
    text = (exclusions or "").strip()
    if not text:
        return []
    entries: list[tuple[str, object]] = [("title", "Scope specifically excluded:")]
    for line in (exclusions or "").split("\n"):
        entries.append(("scope", line))
    return entries


def _build_extra_section_entries(extra_sections, fee_xml_template, prop_fee_xml_template):
    """Build (kind, value) entries for additional scope sections in the Word doc."""
    entries: list[tuple[str, object]] = []
    for extra in extra_sections:
        # Support 3-tuple (legacy), 5-tuple (with billing info), and 6-tuple
        # (adds exclusions). Older saved payloads roundtrip through the
        # shorter shapes, so default the missing fields.
        if len(extra) >= 6:
            title, scope, fee, bt, nte_flag, exclusions = extra[:6]
        elif len(extra) == 5:
            title, scope, fee, bt, nte_flag = extra
            exclusions = ""
        else:
            title, scope, fee = extra[:3]
            bt, nte_flag, exclusions = "fixed", False, ""
        if title:
            entries.append(("title", f"{title}:"))
        for line in (scope or "").split("\n"):
            entries.append(("scope", line))
        entries.extend(_exclusions_entries(exclusions))
        # Insert a copy of "PROPOSED FEE:" paragraph.
        if prop_fee_xml_template is not None:
            entries.append(("xml", deepcopy(prop_fee_xml_template)))
        # Insert fee-sentence paragraph with correct billing-mode text.
        if fee_xml_template is not None:
            fee_xml = deepcopy(fee_xml_template)
            full_text = build_fee_text(fee, bt, nte_flag)
            for t_el in fee_xml.iter(qn("w:t")):
                if t_el.text:
                    t_el.text = ""
            first_t = next(fee_xml.iter(qn("w:t")), None)
            if first_t is not None:
                first_t.text = full_text
            entries.append(("xml", fee_xml))
    return entries


# ── main entrypoint ──────────────────────────────────────────────────────────

def generate_proposal(values: dict, output_dir: str,
                      template_path: str = TEMPLATE_CONSULTING,
                      section1_fee: str = "",
                      section1_billing_type: str = "fixed",
                      section1_nte: bool = False,
                      extra_sections: list | None = None,
                      output_filename: str | None = None) -> str:
    """Render the Word template to a new .docx and return the output path.

    extra_sections: list of (title, scope, fee, billing_type, nte) tuples for sections 2+.

    output_filename: optional fixed filename (e.g. ``"Project - Proposal v2.docx"``).
    If provided and the file already exists, it's overwritten. If not provided,
    a collision-suffixed unique path is used (legacy behavior).
    """
    doc = Document(template_path)

    scope_ph       = PLACEHOLDERS["scope_of_work"]
    scope_title_ph = PLACEHOLDERS["scope_title"]
    fee_ph         = PLACEHOLDERS["fee"]

    # Locate key paragraphs BEFORE any replacement.
    scope_para       = next((p for p in doc.paragraphs if scope_ph in p.text),       None)
    scope_title_para = next((p for p in doc.paragraphs if scope_title_ph in p.text), None)
    fee_para         = next((p for p in doc.paragraphs if fee_ph in p.text),          None)
    prop_fee_para    = next((p for p in doc.paragraphs
                             if "PROPOSED FEE" in p.text and fee_ph not in p.text), None)

    # Save XML templates for the PROPOSED FEE block before replacement alters them.
    fee_xml_template      = deepcopy(fee_para._p)      if fee_para      else None
    prop_fee_xml_template = deepcopy(prop_fee_para._p) if prop_fee_para else None

    # Split section-1 scope into lines -- never pass \n into a single run.
    first_scope_lines = (values.get("scope_of_work") or "").split("\n")
    replacements = {ph: values.get(key, "") for key, ph in PLACEHOLDERS.items()}
    replacements[scope_ph] = first_scope_lines[0]
    replacements[fee_ph]   = format_fee_for_doc(section1_fee)

    for para in doc.paragraphs:
        for old, new in replacements.items():
            _replace_in_paragraph(para, old, new)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old, new in replacements.items():
                        _replace_in_paragraph(para, old, new)

    # Header text-box placeholders use different casing from body placeholders.
    header_replacements = {
        "{{Client_Name}}":  values.get("client_name",  ""),
        "{{Project_Name}}": values.get("project_name", ""),
        "{{Date}}":         values.get("date",          ""),
    }

    for section in doc.sections:
        for hdr in (section.header, section.footer):
            if hdr:
                for para in hdr.paragraphs:
                    for old, new in replacements.items():
                        _replace_in_paragraph(para, old, new)
                _replace_in_header_textboxes(hdr._element, header_replacements)

    # Overwrite the fee paragraph with the correct billing-mode text.
    if fee_para is not None:
        full_fee_text = build_fee_text(section1_fee, section1_billing_type, section1_nte)
        if fee_para.runs:
            fee_para.runs[0].text = full_fee_text
            for run in fee_para.runs[1:]:
                run.text = ""

    # Phase 1: insert remaining scope lines for section 1 after scope_para,
    # then the section-1 Exclusions block (if any) immediately under it.
    if scope_para is not None and scope_title_para is not None:
        s1_entries: list[tuple[str, object]] = [
            ("scope", line) for line in first_scope_lines[1:]
        ]
        s1_entries.extend(_exclusions_entries(values.get("scope_excluded", "")))
        if s1_entries:
            _insert_scope_paragraphs(scope_para, scope_title_para, s1_entries)

    # Phase 2: insert additional sections after the fee paragraph.
    if fee_para is not None and scope_title_para is not None and extra_sections:
        extra_entries = _build_extra_section_entries(
            extra_sections, fee_xml_template, prop_fee_xml_template)
        if extra_entries:
            _insert_scope_paragraphs(fee_para, scope_title_para, extra_entries)

    if output_filename:
        out_path = os.path.join(output_dir, output_filename)
        # Overwrite-in-place is what we want for hash-aware versioned naming.
        if os.path.exists(out_path):
            try:
                os.remove(out_path)
            except OSError:
                # Another process is holding the file (e.g. Word has it open).
                # Fall back to a collision-suffixed name so generation never fails.
                stem, ext = os.path.splitext(output_filename)
                out_path = unique_path(output_dir, stem, ext.lstrip("."))
    else:
        out_path = unique_path(
            output_dir, f"{safe_name(values['project_name'])} - Proposal", "docx"
        )
    doc.save(out_path)
    return out_path
